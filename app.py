from flask import Flask, render_template, request, redirect, url_for, flash, session, send_file
from pymongo import MongoClient
from werkzeug.security import generate_password_hash, check_password_hash
from bson.objectid import ObjectId
import uuid
import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import matplotlib.pyplot as plt
from io import BytesIO

app = Flask(__name__)
app.secret_key = "your_secret_key"

# MongoDB connection
client = MongoClient("mongodb+srv://venkat42005:djWfghShxTH464Pr@cluster0.lvdaz.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0")
db = client["feedback_form_db"]
forms_collection = db["forms"]
feedback_collection = db["feedback"]
users_collection = db["users"]
departments_collection = db["departments"]

# Ensure default admin account exists
if not users_collection.find_one({"username": "principal"}):
    users_collection.insert_one({
        "username": "principal",
        "password": generate_password_hash("principal123"),
        "role": "admin"
    })

# Ensure default departments exist
initial_departments = [
    {"name": "Artificial Intelligence & Data Science", "abbr": "ai_ds"},
    {"name": "Civil Engineering", "abbr": "civil"},
    {"name": "Computer Science & Engineering", "abbr": "cse"},
    {"name": "Computer Science and Engineering (Artificial Intelligence & Machine Learning)", "abbr": "cse_aiml"},
    {"name": "Computer Science and Engineering (Cyber Security)", "abbr": "cse_cs"},
    {"name": "Electrical & Electronics Engineering", "abbr": "eee"},
    {"name": "Electronics & Communication Engineering", "abbr": "ece"},
    {"name": "Electronics Engineering (VLSI Design and Technology)", "abbr": "vlsi"},
    {"name": "Information Technology", "abbr": "it"},
    {"name": "Mechanical Engineering", "abbr": "mech"},
    {"name": "Mechatronics Engineering", "abbr": "mct"},
    {"name": "Science & Humanities", "abbr": "sci_hum"},
    {"name": "Master of Business Administration", "abbr": "mba"}
]

# Insert each department if it does not already exist
for dept in initial_departments:
    if not departments_collection.find_one({"abbr": dept["abbr"]}):
        departments_collection.insert_one(dept)


# Ensure HOD accounts exist for departments
for dept in departments_collection.find():
    username = f"hod_{dept['abbr']}"
    if not users_collection.find_one({"username": username}):
        users_collection.insert_one({
            "username": username,
            "password": generate_password_hash(f"{username}123"),
            "role": "hod",
            "department": dept["name"]
        })

# Routes
@app.route('/')
def home():
    if 'role' in session:
        if session['role'] == 'admin':
            return redirect(url_for('admin_dashboard'))
        elif session['role'] == 'hod':
            return redirect(url_for('hod_dashboard'))
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username'].strip()
        password = request.form['password'].strip()
        user = users_collection.find_one({"username": username})

        if user and check_password_hash(user['password'], password):
            session['user_id'] = str(user['_id'])
            session['username'] = user['username']
            session['role'] = user['role']
            if user['role'] == 'hod':
                session['department'] = user.get('department', '')

            flash(f"Welcome, {user['role'].capitalize()}!", "success")
            return redirect(url_for(f"{user['role']}_dashboard"))
        else:
            flash("Invalid username or password", "danger")
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    flash("Logged out successfully.", "success")
    return redirect(url_for('login'))

@app.route('/manage_departments', methods=['GET', 'POST'])
def manage_departments():
    """
    Manage departments (Add, Edit, Delete) functionality for the admin.
    """
    if 'role' not in session or session['role'] != 'admin':
        flash("Unauthorized access.", "danger")
        return redirect(url_for('login'))

    try:
        if request.method == 'POST':
            action = request.form.get('action', '').strip()

            if action == 'add':
                dept_name = request.form.get('dept_name', '').strip()
                dept_abbr = request.form.get('dept_abbr', '').strip().lower()

                if not dept_name or not dept_abbr:
                    flash("Both Department Name and Abbreviation are required.", "danger")
                elif departments_collection.find_one({"abbr": dept_abbr}):
                    flash("Department abbreviation already exists.", "danger")
                else:
                    # Add new department
                    departments_collection.insert_one({"name": dept_name, "abbr": dept_abbr})
                    flash("Department added successfully.", "success")

            elif action == 'edit':
                dept_id = request.form.get('dept_id', '').strip()
                new_name = request.form.get('new_name', '').strip()
                new_abbr = request.form.get('new_abbr', '').strip().lower()

                if not dept_id or not new_name or not new_abbr:
                    flash("All fields are required for editing.", "danger")
                else:
                    # Update the department details
                    result = departments_collection.update_one(
                        {"_id": ObjectId(dept_id)},
                        {"$set": {"name": new_name, "abbr": new_abbr}}
                    )
                    if result.matched_count:
                        flash("Department updated successfully.", "success")
                    else:
                        flash("Department not found.", "danger")

            elif action == 'delete':
                dept_id = request.form.get('dept_id', '').strip()

                if not dept_id:
                    flash("Department ID is required to delete.", "danger")
                else:
                    # Delete the department
                    result = departments_collection.delete_one({"_id": ObjectId(dept_id)})
                    if result.deleted_count:
                        flash("Department deleted successfully.", "success")
                    else:
                        flash("Department not found.", "danger")

        # Fetch all departments for rendering
        departments = list(departments_collection.find())
        return render_template('manage_departments.html', departments=departments)

    except Exception as e:
        flash(f"An error occurred: {str(e)}", "danger")
        return redirect(url_for('manage_departments'))

@app.route('/manage_users', methods=['GET', 'POST'])
def manage_users():
    """
    Manage HOD users and allow admin to update their own credentials.
    """
    if 'role' not in session or session['role'] != 'admin':
        flash("Unauthorized access.", "danger")
        return redirect(url_for('login'))

    try:
        if request.method == 'POST':
            action = request.form.get('action', '').strip()

            if action == 'add':
                # Add a new HOD user
                username = request.form.get('new_username', '').strip()
                department_name = request.form.get('department', '').strip()
                password = f"{username}123"  # Default password

                if not username or not department_name:
                    flash("Both Username and Department are required.", "danger")
                elif users_collection.find_one({"username": username}):
                    flash("Username already exists. Please choose a different username.", "danger")
                else:
                    users_collection.insert_one({
                        "username": username,
                        "password": generate_password_hash(password),
                        "role": "hod",
                        "department": department_name
                    })
                    flash(f"HOD {username} added successfully with default password.", "success")

            elif action == 'update_credentials':
                # Update credentials for an existing HOD
                hod_id = request.form.get('hod_id', '').strip()
                new_username = request.form.get('new_username', '').strip()
                new_password = request.form.get('new_password', '').strip()

                if not hod_id or not new_username or not new_password:
                    flash("All fields are required to update credentials.", "danger")
                else:
                    user = users_collection.find_one({"_id": ObjectId(hod_id)})
                    if user and user['role'] == 'hod':
                        users_collection.update_one(
                            {"_id": ObjectId(hod_id)},
                            {"$set": {
                                "username": new_username,
                                "password": generate_password_hash(new_password)
                            }}
                        )
                        flash(f"Credentials updated successfully for {new_username}.", "success")
                    else:
                        flash("User not found or is not an HOD.", "danger")

            elif action == 'reset_hod_credentials':
                # Reset HOD credentials to default
                hod_id = request.form.get('hod_id', '').strip()

                if not hod_id:
                    flash("HOD ID is required to reset credentials.", "danger")
                else:
                    user = users_collection.find_one({"_id": ObjectId(hod_id)})
                    if user and user['role'] == 'hod':
                        default_password = f"{user['username']}123"
                        users_collection.update_one(
                            {"_id": ObjectId(hod_id)},
                            {"$set": {"password": generate_password_hash(default_password)}}
                        )
                        flash(f"Password reset successfully for {user['username']} to default.", "success")
                    else:
                        flash("User not found or is not an HOD.", "danger")

            elif action == 'delete':
                # Delete an HOD user
                hod_id = request.form.get('hod_id', '').strip()

                if not hod_id:
                    flash("HOD ID is required to delete the user.", "danger")
                else:
                    result = users_collection.delete_one({"_id": ObjectId(hod_id)})
                    if result.deleted_count > 0:
                        flash("User deleted successfully.", "success")
                    else:
                        flash("User not found. No deletion occurred.", "danger")

            elif action == 'update_admin_credentials':
                # Update admin credentials
                new_admin_username = request.form.get('admin_username', '').strip()
                new_admin_password = request.form.get('admin_password', '').strip()

                if not new_admin_username or not new_admin_password:
                    flash("Both Admin Username and Password are required.", "danger")
                else:
                    admin_user = users_collection.find_one({"role": "admin"})
                    if admin_user:
                        users_collection.update_one(
                            {"_id": admin_user["_id"]},
                            {"$set": {
                                "username": new_admin_username,
                                "password": generate_password_hash(new_admin_password)
                            }}
                        )
                        flash("Admin credentials updated successfully!", "success")
                        session['username'] = new_admin_username
                    else:
                        flash("Admin user not found.", "danger")

        # Fetch users and departments to display
        users = list(users_collection.find({"role": "hod"}))
        departments = list(departments_collection.find({}, {"_id": 0, "name": 1}))

        return render_template('manage_users.html', users=users, departments=departments)

    except Exception as e:
        flash(f"An error occurred: {str(e)}", "danger")
        return redirect(url_for('manage_users'))

@app.route('/some_protected_route')
def protected_route():
    if 'role' not in session:
        # Flash only the "Unauthorized access" message
        flash("Unauthorized access", "danger")
        return redirect(url_for('login'))



# Admin dashboard (accessible by Principal)
@app.route('/admin_dashboard')
def admin_dashboard():
    """
    Admin Dashboard: Displays forms and allows filtering, pagination, and management actions.
    """
    if 'role' not in session or session['role'] != 'admin':
        flash("Unauthorized access.", "danger")
        return redirect(url_for('login'))

    # Get filters from query parameters
    academic_year = request.args.get('academicYear', '').strip()
    department = request.args.get('department', '').strip()
    semester = request.args.get('semester', '').strip()
    batch = request.args.get('batch', '').strip()
    page = int(request.args.get('page', 1))
    per_page = 10  # Default number of forms per page

    # Build the query for filtering
    query = {}
    if academic_year:
        query['academic_year'] = {"$regex": academic_year, "$options": "i"}
    if department:
        query['department'] = {"$regex": department, "$options": "i"}
    if semester:
        query['semester'] = semester
    if batch:
        query['batch'] = {"$regex": batch, "$options": "i"}

    # Fetch and filter forms
    filtered_forms = list(forms_collection.find(query).sort("academic_year", 1))  # Sort by academic year
    total_forms = len(filtered_forms)

    is_filtered = any([academic_year, department, semester, batch])

    # Apply pagination only if no filters are applied
    if not is_filtered:
        total_pages = (total_forms + per_page - 1) // per_page  # Calculate total pages
        start = (page - 1) * per_page
        end = start + per_page
        paginated_forms = filtered_forms[start:end]
    else:
        paginated_forms = filtered_forms
        total_pages = 1

    # Enrich form data for rendering
    for idx, form in enumerate(paginated_forms, start=(page - 1) * per_page + 1):
        form['serial_no'] = idx
        semester_value = int(form.get("semester", 0))
        if semester_value in [1, 2]:
            form["year"] = "First Year"
        elif semester_value in [3, 4]:
            form["year"] = "Second Year"
        elif semester_value in [5, 6]:
            form["year"] = "Third Year"
        elif semester_value in [7, 8]:
            form["year"] = "Fourth Year"
        else:
            form["year"] = "Unknown Year"

    # Flash a message if no forms are found
    if not paginated_forms and not is_filtered:
        flash("No forms found.", "info")

    # Fetch all departments from the database for the dropdown
    departments = list(db['departments'].find({}, {"_id": 0, "name": 1}))

    # Render the admin dashboard template
    return render_template(
        'admin_dashboard.html',
        forms=paginated_forms,
        total_pages=total_pages,
        current_page=page,
        total_forms=total_forms,
        filters={
            'academicYear': academic_year,
            'department': department,
            'semester': semester,
            'batch': batch
        },
        is_filtered=is_filtered,
        departments=departments,
        str=str
    )


# HOD dashboard
@app.route('/hod_dashboard')
def hod_dashboard():
    if 'role' not in session or session['role'] != 'hod':
        flash("Unauthorized access.", "danger")
        return redirect(url_for('login'))

    # Retrieve department from the session
    department = session.get('department', '')
    if not department:
        flash("Your account is not associated with any department.", "danger")
        return redirect(url_for('logout'))

    # Filters
    academic_year = request.args.get('academicYear', '').strip()
    semester = request.args.get('semester', '').strip()
    batch = request.args.get('batch', '').strip()
    page = int(request.args.get('page', 1))
    per_page = 10

    # Query forms related to the logged-in HOD's department
    query = {'department': department}
    if academic_year:
        query['academic_year'] = {"$regex": academic_year, "$options": "i"}
    if semester:
        query['semester'] = semester
    if batch:
        query['batch'] = {"$regex": batch, "$options": "i"}

    # Fetch filtered forms
    filtered_forms = list(forms_collection.find(query).sort("academic_year", 1))
    total_forms = len(filtered_forms)

    is_filtered = any([academic_year, semester, batch])

    # Apply pagination if no filters are applied
    if not is_filtered:
        total_pages = (total_forms + per_page - 1) // per_page
        start = (page - 1) * per_page
        end = start + per_page
        paginated_forms = filtered_forms[start:end]
    else:
        paginated_forms = filtered_forms
        total_pages = 1

    for idx, form in enumerate(paginated_forms, start=(page - 1) * per_page + 1):
        form['serial_no'] = idx
        semester_value = int(form.get("semester", 0))
        if semester_value in [1, 2]:
            form["year"] = "First Year"
        elif semester_value in [3, 4]:
            form["year"] = "Second Year"
        elif semester_value in [5, 6]:
            form["year"] = "Third Year"
        elif semester_value in [7, 8]:
            form["year"] = "Fourth Year"
        else:
            form["year"] = "Unknown Year"

    return render_template(
        'hod_dashboard.html',
        forms=paginated_forms,
        total_pages=total_pages,
        current_page=page,
        total_forms=total_forms,
        department=department,  # Pass department to the template
        filters={
            'academicYear': academic_year,
            'semester': semester,
            'batch': batch
        },
        is_filtered=is_filtered,
        str=str
    )


# Feedback Form (accessible by anyone with the form link)
@app.route('/feedback_form/<form_id>', methods=['GET', 'POST'])
def feedback_form(form_id):
    form_details = forms_collection.find_one({"_id": form_id})
    if not form_details:
        flash("Feedback form not found.", "danger")
        return "Form not found", 404

    # Debug to ensure labs are present
    print(form_details)  # Check the structure of form_details

    semester = int(form_details.get("semester", 1))
    semester_type = "Odd Semester" if semester % 2 != 0 else "Even Semester"

    return render_template(
        'feedback_form.html',
        form_details=form_details,
        semester_type=semester_type
    )

# Submit Feedback
@app.route('/submit_feedback', methods=['POST'])
def submit_feedback():
    try:
        # Retrieve the form ID from the submitted form
        form_id = request.form.get('form_id')
        if not form_id:
            flash("Form ID is missing!", "danger")
            return redirect(request.referrer or url_for('home'))

        # Fetch the form details from the database
        form_details = forms_collection.find_one({"_id": form_id})
        if not form_details:
            flash("Feedback form not found.", "danger")
            return redirect(request.referrer or url_for('home'))

        # Prepare feedback data for theory courses
        feedback_data = []

        # Handle course feedback
        for course in form_details.get("courses", []):
            course_code = course["course_code"]
            course_feedback = {
                "course_code": course_code,
                "feedback": {}
            }
            for i in range(1, 9):  # Iterate through 8 theory questions
                question_key = f"q{i}_{course_code}"
                course_feedback["feedback"][f"q{i}"] = request.form.get(question_key)
            course_feedback["suggestions"] = request.form.get(f"suggestions_{course_code}")
            feedback_data.append(course_feedback)

        # Handle lab feedback
        lab_feedback_data = []
        for lab in form_details.get("labs", []):
            lab_code = lab["lab_code"]
            lab_feedback = {
                "lab_code": lab_code,
                "feedback": {}
            }
            for i in range(1, 5):  # Iterate through 4 lab-specific questions
                question_key = f"lab_q{i}_{lab_code}"
                lab_feedback["feedback"][f"lab_q{i}"] = request.form.get(question_key)
            lab_feedback["suggestions"] = request.form.get(f"lab_suggestions_{lab_code}")
            lab_feedback_data.append(lab_feedback)

        # Prepare the document for insertion into the feedback collection
        feedback_document = {
            "form_id": form_id,
            "course_feedback_data": feedback_data,
            "lab_feedback_data": lab_feedback_data,
        }

        # Insert the feedback into the database
        feedback_collection.insert_one(feedback_document)

        # Redirect to the "Thank You" page
        return redirect(url_for('thank_you', form_id=form_id))

    except Exception as e:
        # Log the exception for debugging
        print(f"Error submitting feedback: {e}")

        # Display a friendly error message and redirect back to the form
        flash("An error occurred while submitting your feedback. Please try again.", "danger")
        return redirect(request.referrer or url_for('home'))

@app.route('/thank_you/<form_id>')
def thank_you(form_id):
    return render_template('thank_you.html', form_id=form_id)


@app.route('/view_report/<form_id>')
def view_report(form_id):
    if 'role' not in session or session['role'] != 'admin':
        flash("Unauthorized access", "danger")
        return redirect(url_for('login'))

    form_details = forms_collection.find_one({"_id": form_id})
    if not form_details:
        return "Form not found", 404

    header_image = url_for('static', filename='images/header.jpg')

    semester = int(form_details.get("semester", 1))
    semester_type = "Odd Semester" if semester % 2 != 0 else "Even Semester"

    feedback_data = list(feedback_collection.find({"form_id": form_id}))

    student_count = int(form_details.get("students_strength", 0))
    students_participated = len(feedback_data)

    # Process theory courses
    course_averages = {}
    faculty_reports = []
    for course in form_details.get("courses", []):
        course_code = course["course_code"]
        course_title = course["course_name"]
        staff_name = course["staff_name"]

        ratings = []
        question_means = {f"q{i}": [] for i in range(1, 9)}
        suggestions = []

        for feedback in feedback_data:
            for data in feedback.get("course_feedback_data", []):
                if data["course_code"] == course_code:
                    ratings.extend(int(data["feedback"].get(f"q{i}", 0)) for i in range(1, 9))
                    for i in range(1, 9):
                        question_means[f"q{i}"].append(int(data["feedback"].get(f"q{i}", 0)))
                    if data.get("suggestions"):
                        suggestions.append(data["suggestions"])

        course_averages[course_code] = round(sum(ratings) / len(ratings), 2) if ratings else 0

        question_means_avg = [
            (i + 1, round(sum(values) / len(values), 2) if values else 0)
            for i, values in enumerate(question_means.values())
        ]

        faculty_reports.append({
            "course_code": course_code,
            "course_title": course_title,
            "staff_name": staff_name,
            "question_means": question_means_avg,
            "suggestions": suggestions,
        })

    # Process lab courses
    lab_averages = {}
    lab_reports = []
    for lab in form_details.get("labs", []):
        lab_code = lab["lab_code"]
        lab_name = lab["lab_name"]
        lab_instructors = lab["lab_instructors"]

        ratings = []
        question_means = {f"lab_q{i}": [] for i in range(1, 5)}
        suggestions = []

        for feedback in feedback_data:
            for data in feedback.get("lab_feedback_data", []):
                if data["lab_code"] == lab_code:
                    ratings.extend(int(data["feedback"].get(f"lab_q{i}", 0)) for i in range(1, 5))
                    for i in range(1, 5):
                        question_means[f"lab_q{i}"].append(int(data["feedback"].get(f"lab_q{i}", 0)))
                    if data.get("suggestions"):
                        suggestions.append(data["suggestions"])

        lab_averages[lab_code] = round(sum(ratings) / len(ratings), 2) if ratings else 0

        question_means_avg = [
            (i + 1, round(sum(values) / len(values), 2) if values else 0)
            for i, values in enumerate(question_means.values())
        ]

        lab_reports.append({
            "lab_code": lab_code,
            "lab_name": lab_name,
            "lab_instructors": lab_instructors,
            "question_means": question_means_avg,
            "suggestions": suggestions,
        })

    return render_template(
        'report.html',
        header_image=header_image,
        form_details=form_details,
        student_count=student_count,
        students_participated=students_participated,
        course_averages=course_averages,
        faculty_reports=faculty_reports,
        lab_averages=lab_averages,
        lab_reports=lab_reports,
        semester_type=semester_type
    )

@app.route('/create_form', methods=['GET', 'POST'])
def create_form():
    """
    Handles creating a new feedback form.
    Accessible by users with 'admin' or 'hod' roles.
    """
    if 'role' not in session or session['role'] not in ['hod', 'admin']:
        flash("Unauthorized access", "danger")
        return redirect(url_for('login'))

    # Determine the department based on the user's role
    department = session.get('department', '') if session['role'] == 'hod' else None
    departments = (
        list(departments_collection.find({}, {"_id": 0, "name": 1})) if session['role'] == 'admin' else []
    )

    if request.method == 'POST':
        try:
            # Debugging - Print submitted form data
            print("Form Data Received:", request.form.to_dict())

            academic_year = request.form['academicYear']
            semester = request.form['semester']
            batch = request.form['batch']
            students_strength = int(request.form['studentsStrength'])

            # Ensure the department is assigned based on the role
            if session['role'] == 'hod':
                department = session.get('department', '')
            else:
                department = request.form['department']

            # Validate that department is assigned
            if not department:
                flash("Department not assigned. Cannot create form.", "danger")
                return redirect(url_for('create_form'))

            # Prepare course data
            courses = []
            course_count = int(request.form['courseCount'])
            for i in range(1, course_count + 1):
                course_code = request.form[f'courseCode{i}']
                course_title = request.form[f'courseTitle{i}']
                staff_name = request.form[f'staffName{i}']
                courses.append({
                    'course_code': course_code,
                    'course_name': course_title,
                    'staff_name': staff_name,
                })

            # Prepare lab data
            labs = []
            lab_count = int(request.form['labCount'])
            for i in range(1, lab_count + 1):
                lab_code = request.form.get(f'labCode{i}', '').strip()
                lab_title = request.form.get(f'labTitle{i}', '').strip()
                lab_instructors = request.form.get(f'labInstructors{i}', '').strip()
                if lab_code and lab_title and lab_instructors:
                    labs.append({
                        'lab_code': lab_code,
                        'lab_name': lab_title,
                        'lab_instructors': lab_instructors,
                    })

            # Insert the new form into the database
            form_id = str(uuid.uuid4())
            new_form = {
                '_id': form_id,
                'academic_year': academic_year,
                'department': department,
                'semester': semester,
                'batch': batch,
                'students_strength': students_strength,
                'courses': courses,
                'labs': labs,
            }
            print("Form to Insert:", new_form)  # Debugging - Log the form data
            forms_collection.insert_one(new_form)
            flash("Form created successfully!", "success")

            # Redirect to the appropriate dashboard
            return redirect(url_for('admin_dashboard') if session['role'] == 'admin' else url_for('hod_dashboard'))

        except Exception as e:
            print(f"Error creating form: {e}")  # Debugging - Log the error
            flash(f"Error creating form: {str(e)}", "danger")

    return render_template(
        'create_form.html',
        is_edit=False,
        form=None,
        departments=departments,
        department=department,
        role=session['role'],
    )

@app.route('/edit_form/<form_id>', methods=['GET', 'POST'])
def edit_form(form_id):
    """
    Handles editing an existing feedback form.
    Accessible by users with 'admin' or 'hod' roles.
    """
    if 'role' not in session or session['role'] not in ['admin', 'hod']:
        flash("Unauthorized access.", "danger")
        return redirect(url_for('login'))

    # Fetch the form to edit
    form = forms_collection.find_one({"_id": form_id})
    if not form:
        flash("Form not found.", "danger")
        return redirect(url_for('hod_dashboard') if session['role'] == 'hod' else url_for('admin_dashboard'))

    # Fetch department dropdown for admin
    departments = (
        list(departments_collection.find({}, {"_id": 0, "name": 1})) if session['role'] == 'admin' else []
    )

    if request.method == 'POST':
        try:
            # Debugging - Print submitted form data
            print("Form Data Received for Update:", request.form.to_dict())

            academic_year = request.form['academicYear']
            semester = request.form['semester']
            batch = request.form['batch']
            students_strength = int(request.form['studentsStrength'])

            # Update the department based on the role
            if session['role'] == 'admin':
                department = request.form['department']
            else:
                department = session.get('department', '')

            # Validate the department
            if not department:
                flash("Department is required.", "danger")
                return redirect(url_for('edit_form', form_id=form_id))

            # Prepare courses data
            courses = []
            course_count = int(request.form['courseCount'])
            for i in range(1, course_count + 1):
                course_code = request.form[f'courseCode{i}']
                course_title = request.form[f'courseTitle{i}']
                staff_name = request.form[f'staffName{i}']
                courses.append({
                    'course_code': course_code,
                    'course_name': course_title,
                    'staff_name': staff_name,
                })

            # Prepare labs data
            labs = []
            lab_count = int(request.form['labCount'])
            for i in range(1, lab_count + 1):
                lab_code = request.form.get(f'labCode{i}', '').strip()
                lab_title = request.form.get(f'labTitle{i}', '').strip()
                lab_instructors = request.form.get(f'labInstructors{i}', '').strip()
                if lab_code and lab_title and lab_instructors:
                    labs.append({
                        'lab_code': lab_code,
                        'lab_name': lab_title,
                        'lab_instructors': lab_instructors,
                    })

            # Update the form in the database
            update_data = {
                "academic_year": academic_year,
                "department": department,
                "semester": semester,
                "batch": batch,
                "students_strength": students_strength,
                "courses": courses,
                "labs": labs,
            }
            print("Form to Update:", update_data)  # Debugging - Log the update data
            forms_collection.update_one({"_id": form_id}, {"$set": update_data})
            flash("Form updated successfully!", "success")

            # Redirect to the appropriate dashboard
            return redirect(url_for('admin_dashboard') if session['role'] == 'admin' else url_for('hod_dashboard'))

        except Exception as e:
            print(f"Error updating form: {e}")  # Debugging - Log the error
            flash(f"Error updating form: {str(e)}", "danger")

    return render_template(
        'create_form.html',
        str=str,
        is_edit=True,
        form=form,
        departments=departments,
        department=form.get('department', ''),
        role=session['role'],
    )


@app.route('/delete_form/<form_id>', methods=['POST'])
def delete_form(form_id):
    """
    Handles deleting a feedback form and its associated feedback data.
    Accessible by users with 'admin' or 'hod' roles.
    """
    if 'role' not in session or session['role'] not in ['admin', 'hod']:
        flash("Unauthorized access.", "danger")
        return redirect(url_for('login'))

    current_page = request.args.get('page', 1)

    try:
        # Attempt to delete the form
        form_result = forms_collection.delete_one({"_id": form_id})
        if form_result.deleted_count > 0:
            # Delete associated feedback data
            feedback_result = feedback_collection.delete_many({"form_id": form_id})
            flash(
                f"Form and its {feedback_result.deleted_count} associated feedback entries deleted successfully!",
                "success",
            )
        else:
            flash("Form not found. No deletion occurred.", "danger")
    except Exception as e:
        flash(f"Error deleting form: {str(e)}", "danger")

    # Redirect to the correct dashboard
    return redirect(url_for('hod_dashboard', page=current_page) if session['role'] == 'hod' else url_for('admin_dashboard', page=current_page))

@app.route('/download_report/<form_id>')
def download_report(form_id):
    """Generate a Word document for the report."""
    form_details = forms_collection.find_one({"_id": form_id})
    if not form_details:
        flash("Form not found.", "danger")
        return redirect(url_for('admin_dashboard'))

    semester = int(form_details.get("semester", 1))
    semester_type = "Odd Semester" if semester % 2 != 0 else "Even Semester"

    # Fetch students strength from form details
    students_strength = form_details.get("students_strength", 0)  # Ensure a default value
    feedback_data = list(feedback_collection.find({"form_id": form_id}))
    students_participated = len(feedback_data)

    course_averages = {}
    faculty_reports = []
    for course in form_details.get("courses", []):
        course_code = course["course_code"]
        course_title = course["course_name"]
        staff_name = course["staff_name"]

        ratings = []
        question_means = {f"q{i}": [] for i in range(1, 9)}
        suggestions = []

        for feedback in feedback_data:
            for data in feedback.get("course_feedback_data", []):
                if data["course_code"] == course_code:
                    ratings.extend(
                        int(data["feedback"].get(f"q{i}", 0)) for i in range(1, 9)
                    )
                    for i in range(1, 9):
                        question_means[f"q{i}"].append(int(data["feedback"].get(f"q{i}", 0)))
                    if data.get("suggestions"):
                        suggestions.append(data["suggestions"])

        course_averages[course_code] = round(sum(ratings) / len(ratings), 2) if ratings else 0

        question_means_avg = [
            (i + 1, round(sum(values) / len(values), 2) if values else 0)
            for i, values in enumerate(question_means.values())
        ]

        faculty_reports.append({
            "course_code": course_code,
            "course_title": course_title,
            "staff_name": staff_name,
            "question_means": question_means_avg,
            "suggestions": suggestions,
        })

    # Process lab reports
    lab_averages = {}
    lab_reports = []
    for lab in form_details.get("labs", []):
        lab_code = lab["lab_code"]
        lab_name = lab["lab_name"]
        lab_instructors = lab["lab_instructors"]

        ratings = []
        question_means = {f"lab_q{i}": [] for i in range(1, 5)}
        suggestions = []

        for feedback in feedback_data:
            for data in feedback.get("lab_feedback_data", []):
                if data["lab_code"] == lab_code:
                    ratings.extend(int(data["feedback"].get(f"lab_q{i}", 0)) for i in range(1, 5))
                    for i in range(1, 5):
                        question_means[f"lab_q{i}"].append(int(data["feedback"].get(f"lab_q{i}", 0)))
                    if data.get("suggestions"):
                        suggestions.append(data["suggestions"])

        lab_averages[lab_code] = round(sum(ratings) / len(ratings), 2) if ratings else 0

        question_means_avg = [
            (i + 1, round(sum(values) / len(values), 2) if values else 0)
            for i, values in enumerate(question_means.values())
        ]

        lab_reports.append({
            "lab_code": lab_code,
            "lab_name": lab_name,
            "lab_instructors": lab_instructors,
            "question_means": question_means_avg,
            "suggestions": suggestions,
        })

    # Helper function to set font style
    def apply_font_style(paragraph, font_name="Bookman Old Style", font_size=12, bold=False):
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = bold

    # Generate Word Document
    document = Document()

    # Add header image
    section = document.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_run = header_paragraph.add_run()
    header_run.add_picture("static/images/header.jpg", width=Pt(450))  # Adjust width as needed

    # Add report title
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title_paragraph.add_run("FACULTY PERFORMANCE – STUDENT’S FEEDBACK\nSUMMARY REPORT\n")
    apply_font_style(title_paragraph, font_size=12, bold=True)
    title_paragraph.add_run(f"Academic Year {form_details['academic_year']} ({semester_type})\n").bold = True
    apply_font_style(title_paragraph, font_size=12, bold=True)

    # Add details in table format
    table = document.add_table(rows=4, cols=2)  # Adjust rows to include students strength
    table.style = "Table Grid"

    details = [
        ("Department", form_details['department']),
        ("Year and Semester", form_details['batch']),
        ("Students Strength", str(students_strength)),
        ("Number of Students Participated", str(students_participated)),
    ]

    for row, (key, value) in enumerate(details):
        cells = table.rows[row].cells
        cells[0].text = key
        cells[1].text = value
        for cell in cells:
            for paragraph in cell.paragraphs:
                apply_font_style(paragraph)

    document.add_paragraph("\n")  # Add some spacing

    # Overall Feedback Report Table
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Course Code & Title"
    hdr_cells[1].text = "Subject Handling Faculty"
    hdr_cells[2].text = "Average Point"

    # Make the header text bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            apply_font_style(paragraph, bold=True)
    for course in form_details["courses"]:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{course['course_code']} - {course['course_name']}"
        row_cells[1].text = course['staff_name']
        row_cells[2].text = f"{course_averages[course['course_code']]:.2f}"
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                apply_font_style(paragraph)

    document.add_page_break()

    # Individual Faculty Reports
    particulars = [
        "Explicitly spells out the learning objectives of the course, various chapters, and evaluation pattern",
        "Coverage and completion of the syllabus",
        "Course material / class notes given",
        "Gives assignment / homework regularly and monitors them properly",
        "Is punctual to the class and engages for the entire hour",
        "Presents subject matter on the board / PPTs, etc., neatly in a format readable by all",
        "Provides feedback about the progress of the students and motivates them by giving tips and advice",
        "Encourages the students to actively participate in the class activities (through discussions, question answers, brainstorming, etc.)"
    ]

    for faculty in faculty_reports:
        document.add_heading("Individual Faculty Report", level=2)
        document.add_paragraph(f"Name of the Faculty: {faculty['staff_name']}")
        document.add_paragraph(f"Designation / Department: {form_details['department']}")
        document.add_paragraph(f"Course Code & Title: {faculty['course_code']} - {faculty['course_title']}")
        document.add_paragraph(f"Class & Semester: {form_details['batch']}")

        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "S.No."
        hdr_cells[1].text = "Particulars"
        hdr_cells[2].text = "Individual Mean"

        for index, mean in faculty["question_means"]:
            row_cells = table.add_row().cells
            row_cells[0].text = str(index)
            row_cells[1].text = particulars[index - 1]
            row_cells[2].text = f"{mean:.2f}"
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    apply_font_style(paragraph)

        document.add_paragraph("Suggestions:")
        for suggestion in faculty["suggestions"]:
            suggestion_paragraph = document.add_paragraph(f"- {suggestion}")
            apply_font_style(suggestion_paragraph)

        document.add_page_break()

    # Lab Reports Section
    lab_particulars = [
        "Does the faculty explain the theoretical concepts and lab instructions related to the experiments clearly?",
        "Does the faculty evaluate observation and record notebooks upon completion of every lab session?",
        "Does the faculty clarify doubts, solve problems, and encourage active participation during the lab?",
        "Are the lab equipment/systems, and tools in ready-to-use condition?"
    ]

    for lab in lab_reports:
        document.add_heading("Lab Feedback Report", level=2)
        document.add_paragraph(f"Name of the Faculty: {lab['lab_instructors']}")
        document.add_paragraph(f"Designation / Department: {form_details['department']}")
        document.add_paragraph(f"Course Code & Title: {lab['lab_code']} - {lab['lab_name']}")
        document.add_paragraph(f"Class & Semester: {form_details['batch']}")

        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "S.No."
        hdr_cells[1].text = "Particulars"
        hdr_cells[2].text = "Individual Mean"

        for index, mean in lab["question_means"]:
            row_cells = table.add_row().cells
            row_cells[0].text = str(index)
            row_cells[1].text = lab_particulars[index - 1]
            row_cells[2].text = f"{mean:.2f}"
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    apply_font_style(paragraph)

        document.add_paragraph("Suggestions:")
        for suggestion in lab["suggestions"]:
            suggestion_paragraph = document.add_paragraph(f"- {suggestion}")
            apply_font_style(suggestion_paragraph)

        document.add_page_break()

    # Save the document
    file_path = f"feedback_report_{form_id}.docx"
    document.save(file_path)

    return send_file(file_path, as_attachment=True, download_name="Feedback_Report.docx")

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))
